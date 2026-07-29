"""
MIA — Generador automático del reporte mensual
===============================================
Con el dato nuevo ya ensamblado, arma el MODELO de reporte mensual: 3 gráficos
(consolidado, desagregado por eje y núcleo), las tablas comparativas y un .docx con el
DISEÑO INSTITUCIONAL LyP. El documento se construye ABRIENDO la plantilla oficial
`Modelos y Administración/Modelo documento LyP.docx` como base: así hereda el encabezado
(logo + "INFORME DE POLÍTICAS PÚBLICAS"), el pie (banda roja de CONTACTO), los estilos y
las fuentes exactas de la casa. Solo se vacía el cuerpo y se escribe el contenido.

La parte factual (cifras, tablas, lectura de datos) es automática y determinística; la
"Lectura de Libertad y Progreso" queda como slot para redactar en voz institucional
(ver Documentos/MIA — Prompt lectura institucional.md).

Uso:  py 00_Comun/generar_reporte_mensual.py            # último mes de mia_mensual
      py 00_Comun/generar_reporte_mensual.py --mes 2026-07
Requisitos: pip install pandas matplotlib python-docx pyyaml
"""
from __future__ import annotations
import argparse, logging, re, sys
from pathlib import Path
import pandas as pd
import yaml
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "output"
DOCS = BASE / "Documentos"
PLANTILLA = BASE / "Modelos y Administración" / "Modelo documento LyP.docx"

RED = RGBColor(0xC1, 0x12, 0x1C); GRAY = RGBColor(0x40, 0x40, 0x40)
GRAYL = RGBColor(0x76, 0x76, 0x76); GREEN = RGBColor(0x1E, 0x7A, 0x34); REDN = RGBColor(0xB0, 0x20, 0x20)
EJES = ["sub_Ejecutivo", "sub_Legislativo", "sub_Judicial", "sub_Prensa", "sub_Banco Central"]
EJE_NOM = {"sub_Ejecutivo": "Ejecutivo", "sub_Legislativo": "Legislativo", "sub_Judicial": "Judicial",
           "sub_Prensa": "Prensa", "sub_Banco Central": "Banco Central"}
COLORS = {"Ejecutivo": "#1f77b4", "Legislativo": "#ff7f0e", "Judicial": "#2ca02c",
          "Prensa": "#9467bd", "Banco Central": "#8c564b"}
MESES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
         "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
PUBLICAR = "2024-01"

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)-7s | %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("reporte")


def fmt(x, dec=2):
    return "s/d" if pd.isna(x) else f"{x:.{dec}f}".replace(".", ",")


def signo(x, dec=2):
    return ("+" if x >= 0 else "−") + f"{abs(x):.{dec}f}".replace(".", ",")


def mes_label(per: str) -> str:
    y, m = per.split("-"); return f"{MESES[int(m)]} {y}"


# ---------------- gráficos ----------------
def graficos(df: pd.DataFrame, mes: str):
    d = df.copy(); d.index = pd.PeriodIndex(d.index, freq="M").to_timestamp()
    pub = d[d.index >= pd.Period(PUBLICAR, "M").to_timestamp()]
    ROJO, GRISH = "#C0202A", "#404040"
    fig, ax = plt.subplots(figsize=(9, 4.2), dpi=150)
    ax.plot(pub.index, pub["MIA"], color=ROJO, lw=2.4, marker="o", ms=3)
    ax.set_title("Monitor Institucional Argentino (MIA) — consolidado mensual", color=GRISH, fontsize=12, weight="bold")
    ax.set_ylabel("MIA (0–100)", color=GRISH); ax.grid(alpha=.25)
    last = pub["MIA"].iloc[-1]
    ax.annotate(f"{mes_label(mes)}: {fmt(last,1)}", xy=(pub.index[-1], last), xytext=(-4, 8),
                textcoords="offset points", color=ROJO, weight="bold", fontsize=9, ha="right")
    for s in ("top", "right"): ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(OUT / "mia_consolidado.png", bbox_inches="tight"); plt.close(fig)

    fig, ax = plt.subplots(figsize=(9, 4.6), dpi=150)
    for c in EJES:
        ax.plot(pub.index, pub[c], lw=1.9, label=EJE_NOM[c], color=COLORS[EJE_NOM[c]])
    ax.set_title("Sub-índices por eje — mensual", color=GRISH, fontsize=12, weight="bold")
    ax.set_ylabel("Sub-índice (0–100)", color=GRISH); ax.grid(alpha=.25)
    ax.legend(ncol=5, fontsize=8, loc="lower center", bbox_to_anchor=(.5, -.22), frameon=False)
    for s in ("top", "right"): ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(OUT / "mia_5ejes.png", bbox_inches="tight"); plt.close(fig)


def grafico_nucleo():
    f = OUT / "mia_nucleo_mensual.csv"
    if not f.exists(): return False
    n = pd.read_csv(f)
    col = "MIA_nucleo" if "MIA_nucleo" in n.columns else ("ITR_nucleo" if "ITR_nucleo" in n.columns else None)
    if col is None: return False
    n["t"] = pd.PeriodIndex(n["periodo"], freq="M").to_timestamp()
    ROJO, GRISH = "#C0202A", "#404040"
    fig, ax = plt.subplots(figsize=(9, 4.4), dpi=150)
    ini, fin = n["t"].min(), n["t"].max(); corte = pd.Timestamp("2023-12-31")
    ax.axvspan(ini, min(corte, fin), color="#cfe8f3", alpha=.5, lw=0)
    if fin > corte: ax.axvspan(corte, fin, color="#e6dbf2", alpha=.6, lw=0)
    ax.plot(n["t"], n[col], color=ROJO, lw=2.0)
    ax.set_title("MIA Núcleo — serie mensual comparable (2020 →)", color=GRISH, fontsize=12, weight="bold")
    ax.set_ylabel("MIA Núcleo (0–100)", color=GRISH); ax.grid(alpha=.25)
    for s in ("top", "right"): ax.spines[s].set_visible(False)
    fig.tight_layout(); fig.savefig(OUT / "mia_nucleo.png", bbox_inches="tight"); plt.close(fig)
    return True


# ---------------- docx (usa la plantilla LyP como base) ----------------
def styles_map(doc):
    m = {}
    for s in doc.styles:
        try: m[s.name.lower()] = s
        except Exception: pass
    return m


def pick(sm, *names):
    for n in names:
        if n.lower() in sm: return sm[n.lower()]
    return None


def clear_body(doc):
    body = doc.element.body
    for ch in list(body):
        if ch.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(ch)


def replace_runs(p, pattern, repl):
    runs = list(p.runs)
    full = "".join(r.text or "" for r in runs)
    m = re.search(pattern, full)
    if not m: return False
    s, e = m.span(); pos = 0; placed = False
    for r in runs:
        t = r.text or ""; a, b = pos, pos + len(t); pos = b
        if b <= s or a >= e: continue
        pre = t[:s - a] if a < s else ""
        post = t[e - a:] if b > e else ""
        r.text = (pre + repl + post) if not placed else (pre + post)
        placed = placed or True
    return True


def set_header_date(doc, label):
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            replace_runs(p, r"X\s*20\d\d", label)


def _bold_runs(p, text):
    for i, part in enumerate(text.split("**")):
        if part == "": continue
        r = p.add_run(part)
        if i % 2 == 1: r.font.bold = True
    return p


def shade(cell, hexfill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(sh)


def tabla(doc, tstyle, headers, rows, widths_in, delta_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    if tstyle is not None: t.style = tstyle
    t.autofit = False
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].width = Inches(widths_in[j]); shade(hdr[j], "F2DEDE")
        pr = hdr[j].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        rn = pr.add_run(h); rn.font.bold = True; rn.font.color.rgb = RED; rn.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Inches(widths_in[j])
            pr = cells[j].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            rn = pr.add_run(str(val)); rn.font.size = Pt(9); rn.font.bold = (j == 0 or j == delta_col)
            if delta_col is not None and j == delta_col:
                rn.font.color.rgb = GREEN if str(val).startswith("+") else (REDN if str(val).startswith("−") else GRAY)
    return t


def _add_img(doc, path, cw_in):
    doc.add_picture(str(path), width=Inches(min(6.5, cw_in)))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------- main ----------------
def main() -> int:
    ap = argparse.ArgumentParser(description="MIA — generador del reporte mensual (diseño LyP)")
    ap.add_argument("--mes", help="AAAA-MM (default: último de mia_mensual.csv)")
    args = ap.parse_args()

    df = pd.read_csv(OUT / "mia_mensual.csv"); df["periodo"] = df["periodo"].astype(str); df = df.set_index("periodo")
    mes = args.mes or df.index[-1]
    if mes not in df.index:
        log.error("El mes %s no está en mia_mensual.csv", mes); return 1
    idx = list(df.index); prev = idx[idx.index(mes) - 1] if idx.index(mes) > 0 else None
    y, m = map(int, mes.split("-")); aa_key = f"{y-1}-{m:02d}"; aa = aa_key if aa_key in df.index else None

    estado = "provisional"
    if (OUT / "mia_historico.csv").exists():
        h = pd.read_csv(OUT / "mia_historico.csv", dtype={"periodo": str}); r = h[h["periodo"] == mes]
        if not r.empty: estado = str(r["estado"].iloc[0])
    etiqueta = "cierre definitivo" if estado == "cerrado" else "provisional (nowcast)"

    y_ = yaml.safe_load(open(BASE / "00_Comun" / "variables.yaml", encoding="utf-8"))
    var_eje = {v["var"]: v["eje"] for v in y_["variables"]}
    pesos = {"Ejecutivo": .30, "Legislativo": .20, "Judicial": .20, "Prensa": .15, "Banco Central": .15}

    cur = df.loc[mes]; pr = df.loc[prev] if prev else None
    graficos(df, mes); tiene_nucleo = grafico_nucleo()

    eje_rows = [["MIA (total)", "100%", fmt(pr["MIA"]) if pr is not None else "s/d", fmt(cur["MIA"]),
                 signo(cur["MIA"] - pr["MIA"]) if pr is not None else "—"]]
    for c in EJES:
        nom = EJE_NOM[c]
        eje_rows.append([nom, f"{round(pesos[nom]*100)}%", fmt(pr[c]) if pr is not None else "s/d", fmt(cur[c]),
                         signo(cur[c] - pr[c]) if pr is not None else "—"])

    movers = []
    if pr is not None:
        for var in var_eje:
            if var in df.columns:
                a, b = pr.get(var), cur.get(var)
                if pd.notna(a) and pd.notna(b) and abs(b - a) >= 0.05: movers.append((b - a, var, a, b))
    movers.sort(key=lambda x: -abs(x[0])); movers = movers[:8]

    contribs = []
    if pr is not None:
        for c in EJES: contribs.append((pesos[EJE_NOM[c]] * (cur[c] - pr[c]), EJE_NOM[c], cur[c] - pr[c]))
        contribs.sort(key=lambda x: -abs(x[0]))

    # ---- documento sobre la plantilla LyP ----
    if not PLANTILLA.exists():
        log.error("No encuentro la plantilla: %s", PLANTILLA); return 1
    doc = Document(str(PLANTILLA))
    clear_body(doc)
    set_header_date(doc, mes_label(mes))
    sm = styles_map(doc)
    ST_TITLE = pick(sm, "Title", "Título", "Ttulo")
    ST_SUB = pick(sm, "Subtitle", "Subtítulo")
    ST_H1 = pick(sm, "Heading 1", "heading 1")
    ST_TBL = pick(sm, "Table Grid", "Tabla con cuadrícula")

    sec = doc.sections[0]
    cw_in = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0   # EMU -> pulgadas

    # título + subtítulo con estilos de la casa
    tp = doc.add_paragraph(style=ST_TITLE) if ST_TITLE else doc.add_paragraph()
    rt = tp.add_run("Monitor Institucional Argentino (MIA)"); rt.font.color.rgb = RED
    if not ST_TITLE: rt.font.bold = True; rt.font.size = Pt(24)
    subp = doc.add_paragraph(style=ST_SUB) if ST_SUB else doc.add_paragraph()
    subp.add_run(f"Reporte mensual · Edición {mes_label(mes)} — {etiqueta}")
    doc.add_paragraph("Departamento de Políticas Públicas — Fundación Libertad y Progreso")

    def H(t):
        p = doc.add_paragraph(style=ST_H1) if ST_H1 else doc.add_paragraph()
        r = p.add_run(t); r.font.color.rgb = RED
        if not ST_H1: r.font.bold = True; r.font.size = Pt(13)
        return p
    def P(t):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; _bold_runs(p, t); return p

    # resumen
    H("Resumen ejecutivo")
    partes = [f"El Monitor Institucional Argentino registra en {mes_label(mes)} un valor de **{fmt(cur['MIA'])}** sobre 100 ({etiqueta})."]
    if pr is not None:
        d = cur["MIA"] - pr["MIA"]; verbo = "sube" if d > 0.05 else ("baja" if d < -0.05 else "se mantiene")
        partes.append(f"El índice {verbo} **{signo(d)}** puntos respecto de {mes_label(prev)} ({fmt(pr['MIA'])}).")
    if aa is not None:
        partes.append(f"Interanual ({mes_label(aa)}: {fmt(df.loc[aa,'MIA'])}): {signo(cur['MIA']-df.loc[aa,'MIA'])}.")
    if contribs:
        top = contribs[0]; partes.append(f"El eje que más incide es **{top[1]}** ({'al alza' if top[0] > 0 else 'a la baja'}).")
    P(" ".join(partes))

    H("1. MIA consolidado"); _add_img(doc, OUT / "mia_consolidado.png", cw_in)
    H("2. Sub-índices por eje"); _add_img(doc, OUT / "mia_5ejes.png", cw_in)
    if tiene_nucleo:
        H("3. MIA Núcleo (serie comparable)"); _add_img(doc, OUT / "mia_nucleo.png", cw_in)

    H("4. Lectura mes a mes")
    prev_lbl = mes_label(prev) if prev else "—"
    P(f"Índice y sub-índices, con su peso y la variación frente a {prev_lbl}:")
    ab_prev = prev_lbl.split()[0][:3] if prev else "prev"; ab_cur = mes_label(mes).split()[0][:3]
    w = [1.9, 0.8, 1.1, 1.1, 1.0]; sc = cw_in / sum(w); w = [x * sc for x in w]
    tabla(doc, ST_TBL, ["Eje", "Peso", ab_prev, ab_cur, "Var."], eje_rows, w, delta_col=4)
    doc.add_paragraph()
    if movers:
        P("Principales movimientos por variable:")
        mrows = [[v, var_eje[v], fmt(a, 1), fmt(b, 1), signo(b - a, 1)] for (_, v, a, b) in movers]
        w2 = [2.2, 1.3, 0.9, 0.9, 1.0]; sc2 = cw_in / sum(w2); w2 = [x * sc2 for x in w2]
        tabla(doc, ST_TBL, ["Variable", "Eje", ab_prev, ab_cur, "Var."], mrows, w2, delta_col=4)
        doc.add_paragraph()

    if contribs:
        subs = [f"{nom} ({signo(dsub)})" for _, nom, dsub in contribs[:3]]
        P("Lectura de los datos: el movimiento del mes se explica sobre todo por " + ", ".join(subs) +
          (". Los principales cambios de variables figuran en la tabla anterior." if movers else "."))

    H("5. La lectura de Libertad y Progreso")
    ph = doc.add_paragraph(); rr = ph.add_run(
        "[Redactar en voz institucional LyP — ver «Documentos/MIA — Prompt lectura institucional.md». "
        "Es la única sección que se escribe a mano cada mes; el resto del reporte es automático.]")
    rr.font.italic = True; rr.font.color.rgb = GRAYL

    H("Nota metodológica")
    nn = doc.add_paragraph().add_run(
        "El MIA es un índice mensual de 0 a 100, determinístico y auditable, construido con dato duro oficial "
        "(sin IA en el valor publicado). Combina 18 variables en cinco ejes con pesos fijos (Ejecutivo 30%, "
        "Legislativo 20%, Judicial 20%, Prensa 15%, Banco Central 15%), normalizadas por anclaje al ideal y "
        "suavizadas con media móvil de 12 meses. Se publica desde enero de 2024. El mes en curso sale provisional; "
        "los meses cerrados quedan congelados en el histórico maestro.")
    nn.font.italic = True; nn.font.size = Pt(8.5); nn.font.color.rgb = GRAYL

    DOCS.mkdir(exist_ok=True)
    out_doc = DOCS / f"MIA — Reporte Mensual ({mes_label(mes)}).docx"
    doc.save(str(out_doc))
    log.info("Reporte generado (diseño LyP): %s", out_doc.name)
    print(f"\nOK — {out_doc}\n   MIA {mes}: {fmt(cur['MIA'])} ({etiqueta}) | ancho contenido={cw_in:.2f}in | "
          f"gráficos: consolidado/5ejes" + ("/núcleo" if tiene_nucleo else ""))
    return 0


if __name__ == "__main__":
    sys.exit(main())
